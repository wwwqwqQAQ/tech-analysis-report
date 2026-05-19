#!/usr/bin/env python3
"""科技产业分析日报 — AI/芯片/新能源/互联网/航天，产业政策+技术趋势+投融资+竞争格局"""

import json, os, re, subprocess, sys
from datetime import datetime
from collections import defaultdict, Counter

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
import numpy as np

for name in ("Heiti TC", "PingFang SC", "STHeiti", "Arial Unicode MS"):
    try:
        fm.findfont(name, fallback_to_default=False)
        plt.rcParams["font.sans-serif"] = [name, "DejaVu Sans"]
        break
    except Exception:
        continue
plt.rcParams["axes.unicode_minus"] = False

UA = "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36"
REFERER = "https://tech.sina.com.cn/"
CHART_DIR = "/tmp/tech_charts"
os.makedirs(CHART_DIR, exist_ok=True)

C_RED = "#D32F2F"; C_BLUE = "#1976D2"; C_GRAY = "#616161"; C_ORANGE = "#E65100"
C_GREEN = "#2E7D32"; C_PURPLE = "#7B1FA2"; C_TEAL = "#00796B"

def curl(url):
    r = subprocess.run(["curl", "-4", "-sk", "--connect-timeout", "10",
        "-H", f"User-Agent: {UA}", "-H", f"Referer: {REFERER}", url],
        capture_output=True, timeout=20)
    if r.returncode != 0: return None
    raw = r.stdout
    for enc in ("utf-8", "gbk", "gb2312", "gb18030"):
        try: return raw.decode(enc)
        except (UnicodeDecodeError, LookupError): continue
    return raw.decode("utf-8", errors="replace")

def fetch_sina_search(keywords, size=30):
    """Multi-keyword search via Sina search API."""
    results = []
    seen = set()
    for kw in keywords:
        url = f"https://search.sina.com.cn/api/news?q={kw}&size={size}"
        raw = curl(url)
        if not raw: continue
        try:
            data = json.loads(raw)
            for item in data.get("data", {}).get("list", []):
                u = item.get("url", "")
                if u in seen: continue
                seen.add(u)
                title = re.sub(r'<[^>]+>', '', item.get("title", ""))
                results.append({
                    "title": title, "ctime": str(item.get("ctime", "")),
                    "url": u, "intro": item.get("intro", ""),
                })
        except (json.JSONDecodeError, KeyError): continue
    return results

def fetch_sina_news(lid, num=40):
    all_news = []
    for page in range(1, (num // 20) + 3):
        url = f"https://feed.mix.sina.com.cn/api/roll/get?pageid=153&lid={lid}&k=&num=20&page={page}&r=0.1&callback="
        raw = curl(url)
        if not raw: break
        try:
            data = json.loads(raw)
            items = data.get("result", {}).get("data", [])
            if not items: break
            all_news.extend(items)
            if len(all_news) >= num: break
        except json.JSONDecodeError: break
    return all_news[:num]

# ═══════════════════ CLASSIFICATION ═══════════════════

def classify_tech(news_items):
    """Classify tech news into sub-sectors."""
    cats = {"AI/大模型": [], "半导体/芯片": [], "新能源/电动车": [], "互联网/平台": [],
            "航天/硬科技": [], "政策/监管": [], "投融资/创投": [], "其他": []}
    kw_map = {
        "AI/大模型": ["AI", "人工智能", "大模型", "GPT", "LLM", "深度学习", "机器学习",
                      "NLP", "计算机视觉", "生成式", "智能体", "Agent", "OpenAI",
                      "谷歌", "微软", "Meta", "百度文心", "通义", "DeepSeek", "Kimi",
                      "ChatGPT", "Claude", "Gemini", "Llama", "神经网络"],
        "半导体/芯片": ["芯片", "半导体", "光刻", "晶圆", "台积电", "中芯", "英伟达",
                       "AMD", "英特尔", "高通", "ARM", "RISC-V", "GPU", "CPU",
                       "华为", "海思", "封测", "EDA", "硅", "7nm", "5nm", "3nm"],
        "新能源/电动车": ["新能源", "电动车", "电池", "光伏", "风电", "储能", "氢能",
                         "比亚迪", "特斯拉", "宁德时代", "锂电", "钠电", "固态电池",
                         "充电", "换电", "碳", "双碳", "碳中和"],
        "互联网/平台": ["互联网", "平台", "电商", "社交", "短视频", "直播", "搜索",
                       "字节", "腾讯", "阿里", "美团", "拼多多", "京东", "小红书",
                       "TikTok", "WeChat", "APP", "小程序", "算法", "推荐"],
        "航天/硬科技": ["航天", "卫星", "火箭", "SpaceX", "NASA", "北斗", "空间站",
                       "量子", "核聚变", "机器人", "具身智能", "人形机器人", "脑机",
                       "3D打印", "新材料", "超导", "可控核聚变"],
        "政策/监管": ["政策", "监管", "法规", "反垄断", "数据安全", "隐私", "合规",
                     "审查", "许可", "牌照", "工信部", "网信办", "科技部", "标准"],
        "投融资/创投": ["融资", "投资", "IPO", "上市", "估值", "VC", "PE", "创投",
                       "天使", "A轮", "B轮", "C轮", "独角兽", "退出", "并购", "收购"],
    }
    for item in news_items:
        title = item.get("title", "")
        matched = "其他"
        for cat, kws in kw_map.items():
            if any(kw in title for kw in kws):
                matched = cat; break
        if matched == "其他":
            # Try intro for matching
            intro = item.get("intro", "")
            for cat, kws in kw_map.items():
                if any(kw in intro for kw in kws):
                    matched = cat; break
        cats[matched].append(item)
    return cats

# ═══════════════════ ANALYSIS ═══════════════════

def analyze_company_mentions(items):
    """Track which companies are most mentioned."""
    companies = {
        "华为": ["华为", "海思", "鸿蒙", "昇腾", "鲲鹏"],
        "腾讯": ["腾讯", "微信", "WeChat"],
        "阿里巴巴": ["阿里", "阿里巴巴", "通义", "蚂蚁"],
        "字节跳动": ["字节", "抖音", "TikTok", "豆包"],
        "百度": ["百度", "文心", "萝卜快跑"],
        "比亚迪": ["比亚迪", "BYD"],
        "特斯拉": ["特斯拉", "Tesla"],
        "英伟达": ["英伟达", "NVIDIA", "Nvidia"],
        "OpenAI": ["OpenAI"],
        "谷歌": ["谷歌", "Google", "Gemini", "DeepMind"],
        "微软": ["微软", "Microsoft"],
        "苹果": ["苹果", "Apple"],
        "小米": ["小米", "Xiaomi"],
        "宁德时代": ["宁德时代", "CATL"],
        "台积电": ["台积电", "TSMC"],
        "中芯国际": ["中芯国际", "中芯", "SMIC"],
        "Meta": ["Meta", "Facebook"],
        "DeepSeek": ["DeepSeek", "深度求索"],
    }
    all_text = " ".join(i["title"] for i in items)
    mentions = {}
    for name, kws in companies.items():
        cnt = sum(all_text.count(kw) for kw in kws)
        if cnt > 0: mentions[name] = cnt
    return Counter(mentions).most_common(15)

def analyze_tech_trends(cats, company_mentions):
    """Generate tech trend analysis with insights."""
    trends = []

    # AI analysis
    ai_count = len(cats.get("AI/大模型", []))
    if ai_count > 3:
        ai_items = cats["AI/大模型"]
        # Check for key themes
        agent_ai = sum(1 for i in ai_items if "Agent" in i["title"] or "智能体" in i["title"])
        model_war = sum(1 for i in ai_items if any(k in i["title"] for k in
            ["GPT", "Llama", "Gemini", "文心", "通义", "Kimi", "DeepSeek", "豆包", "Claude"]))
        trends.append({
            "title": "AI大模型竞争格局",
            "signal": "加速分化",
            "analysis": (
                f"AI/大模型领域{ai_count}条新闻。"
                f"{'Agent/智能体成为新热点（{}条），AI从对话式交互向任务执行演进。'.format(agent_ai) if agent_ai > 0 else ''}"
                f"基础模型竞争进入'寡头+开源'双轨阶段：头部闭源模型（GPT、Claude、Gemini、文心）在能力上持续突破，"
                f"开源模型（Llama、DeepSeek、通义千问开源版）在成本和可及性上带来降维打击。"
                f"从技术采用生命周期理论看，AI正从'早期采用者'阶段过渡到'早期大众'阶段，"
                f"2026年下半场的竞争焦点将从'谁的模型最强'转向'谁的生态最完整'——"
                f"工具链、开发者社区、企业级部署能力将成为新的护城河。"
            ),
            "horizon": "6-12个月",
        })

    # Semiconductor
    chip_count = len(cats.get("半导体/芯片", []))
    if chip_count > 3:
        trends.append({
            "title": "半导体产业链博弈",
            "signal": "国产替代加速",
            "analysis": (
                f"半导体/芯片领域{chip_count}条新闻。当前全球半导体产业链处于'效率优先→安全优先'的范式转换期。"
                f"美国主导的技术封锁（设备、EDA、先进制程）与中国推动的自主替代并行推进，"
                f"两端都在加速——这使产业链呈现'双轨化'趋势：一条是美国及其盟友的'可信供应链'，"
                f"另一条是中国主导的'自主可控链'。从产业经济学角度看，这种分裂的代价是"
                f"全球芯片成本上升和重复投资，但也催生了巨大的国产替代市场机会。"
                f"关注中芯国际/华虹的先进制程进展、RISC-V生态成熟度、以及EDA工具国产化率三个关键信号。"
            ),
            "horizon": "1-3年",
        })

    # New energy
    ev_count = len(cats.get("新能源/电动车", []))
    if ev_count > 3:
        trends.append({
            "title": "新能源与电动车产业趋势",
            "signal": "从规模竞争到技术竞争",
            "analysis": (
                f"新能源/电动车领域{ev_count}条新闻。产业正从'补贴驱动的规模扩张'进入'技术驱动的差异化竞争'阶段。"
                f"关键变量包括：固态电池量产时间表（决定下一轮洗牌节奏）、"
                f"智能化程度（智驾+座舱成为区分度核心）、海外市场准入（欧盟反补贴关税、美国IRA法案）。"
                f"从波特五力模型看，新能源车行业当前面临：供应商议价力上升（锂/钴/芯片）、"
                f"买方选择增多（车型供给过剩）、新进入者威胁（小米/华为等跨界玩家）、"
                f"替代品压力（氢能/合成燃料）——行业利润率将承压，集中度将进一步提升。"
            ),
            "horizon": "6-18个月",
        })

    # Policy/Reg
    policy_count = len(cats.get("政策/监管", []))
    if policy_count > 2:
        trends.append({
            "title": "科技监管与产业政策",
            "signal": "精准化监管",
            "analysis": (
                f"科技政策/监管领域{policy_count}条新闻。全球科技监管进入'精准化'阶段——"
                f"不再是笼统的'加强监管'或'放松管制'，而是针对特定领域（AI安全、数据跨境、"
                f"算法推荐、反垄断、芯片出口管制）的精细化规则设计。"
                f"中国的监管特色是'发展-安全双目标平衡'：既要通过科技实现产业升级，"
                f"又要确保数据主权和社会稳定。从制度经济学视角看，这种双目标治理面临内在张力——"
                f"过度安全化会拖慢创新速度，过度发展导向则可能积累系统性风险。"
                f"政策走向的领先指标包括：国务院/发改委/工信部的产业规划更新频率、"
                f"科创板IPO审核节奏、以及跨境数据流动试点城市的扩围情况。"
            ),
            "horizon": "3-12个月",
        })

    return trends

def generate_sector_forecasts(cats):
    """Per-sector technology predictions: 6mo / 1-2yr / 3-5yr."""
    forecasts = []

    ai_count = len(cats.get("AI/大模型", []))
    chip_count = len(cats.get("半导体/芯片", []))
    ev_count = len(cats.get("新能源/电动车", []))
    space_count = len(cats.get("航天/硬科技", []))
    internet_count = len(cats.get("互联网/平台", []))
    invest_count = len(cats.get("投融资/创投", []))

    forecasts.append({
        "sector": "AI / 大模型",
        "hype_cycle": "期望膨胀期→泡沫低谷过渡",
        "short": (
            "Agent/智能体将成为2026下半年的主战场——从'能聊'到'能干'的范式转换加速。"
            "开源模型能力持续逼近闭源模型，API调用价格战白热化。"
            "企业端：从'试试看'到'真金白银上生产系统'，对模型可靠性/安全性的要求会倒逼平台化工具链成熟。"
        ),
        "mid": (
            "1-2年内，模型层的'寡头+开源'格局基本定型：3-5家基础模型商+繁荣的开源生态。"
            "价值创造重心从模型层下沉到应用层和基础设施层——"
            "应用层出现AI-Native的杀手级产品（不只是Copilot类的辅助工具），"
            "基础设施层（向量数据库、模型路由、推理优化、LLMOps）进入整合期。"
        ),
        "long": (
            "3-5年：AGI不会到来，但'足够好的AI'会渗透到每一个知识工作流程。"
            "AI对劳动力市场的结构性冲击开始显现——不是替代人，而是重新定义'会做事'的标准。"
            "从Rogers扩散曲线看，AI将在2028年左右跨越'早期大众'进入'晚期大众'阶段，"
            "届时不用AI的企业会像今天不上网的企业一样罕见。"
        ),
        "confidence": "高",
    })

    forecasts.append({
        "sector": "半导体 / 芯片",
        "hype_cycle": "稳步爬升期（国产替代）+ 技术触发期（新材料/新架构）",
        "short": (
            "先进制程（7nm以下）的国产化仍无时间表，但成熟制程（28nm及以上）国产化率已超30%且在加速。"
            "Chiplet/先进封装成为'弯道超车'的现实路径——用封装技术弥补制程劣势。"
            "RISC-V生态建设进入关键窗口期，能否吸引足够多的软件开发者决定其成败。"
        ),
        "mid": (
            "1-2年内，全球半导体产业'双轨化'格局固化：美国主导的'可信供应链' vs 中国主导的'自主可控链'。"
            "两条链的成本差将成为全球电子产品定价的分水岭——"
            "一部分市场愿意为'安全溢价'买单，另一部分市场追求性价比。"
            "AI推理芯片市场爆发——训练芯片是英伟达的天下，但推理芯片格局远未定型。"
        ),
        "long": (
            "3-5年：硅基芯片逼近物理极限（~1nm），新材料（碳纳米管/二维材料）和"
            "新架构（存算一体/光子计算/量子计算）从实验室走向工程化。"
            "芯片产业的竞争维度从'制程先进性'转向'架构创新性'——"
            "这对中国芯片产业而言既是挑战也是机会窗口。"
        ),
        "confidence": "中高",
    })

    forecasts.append({
        "sector": "新能源 / 电动车",
        "hype_cycle": "稳步爬升期（电动车）+ 期望膨胀期（固态电池/氢能）",
        "short": (
            "电动车：价格战继续，行业集中度提升——3-5家头部车企将占据70%+市场份额。"
            "固态电池：2026年下半年到2027年是量产验证的关键期——"
            "宁德时代/丰田/QuantumScape谁先跑通量产工艺，谁就掌握下一轮洗牌的主动权。"
            "光伏/储能：产能过剩导致价格承压，但出海（中东/非洲/拉美）打开新增量空间。"
        ),
        "mid": (
            "1-2年：智能化（智驾+座舱）替代续航成为电动车核心差异化维度。"
            "从Christensen的'性能供给超过市场需求'理论看，当续航和加速已不再是瓶颈，"
            "竞争焦点自然转向用户每天能感知到的东西——软件体验和智能化程度。"
            "氢能在重卡/航运/工业领域的应用场景逐步清晰，但成本仍是最大障碍。"
        ),
        "long": (
            "3-5年：能源结构从'化石能源基荷+新能源补充'逐步转向'新能源基荷+储能调峰'。"
            "电动车的终极形态不是'带电池的汽车'，而是'带轮子的智能终端'——"
            "汽车产业的估值逻辑将从制造业（低PE）向科技服务业（高PE）迁移。"
        ),
        "confidence": "中",
    })

    forecasts.append({
        "sector": "互联网 / 平台经济",
        "hype_cycle": "生产力 plateau（传统互联网）+ 期望膨胀期（AI-Native应用）",
        "short": (
            "传统互联网增长见顶——流量红利消失、监管常态化、用户时长接近天花板。"
            "增长引擎切换为：AI赋能存量业务提效 + 出海（TEMU/TikTok/Shein的全球化）。"
            "AI搜索/AI社交/AI电商是短期内最可能被AI重构的三个互联网产品品类。"
        ),
        "mid": (
            "1-2年：平台经济的估值逻辑正在被重写——"
            "有AI原生能力的平台享受溢价，纯流量分发平台面临折价。"
            "TikTok的地缘政治命运将影响整个中国互联网出海的战略路径。"
        ),
        "long": "3-5年：互联网不再是一个独立'行业'，而成为所有行业的基础设施——就像电力一样隐身于日常。",
        "confidence": "中",
    })

    forecasts.append({
        "sector": "航天 / 硬科技",
        "hype_cycle": "技术触发期→期望膨胀期过渡",
        "short": (
            "商业航天：中国版'星链'加速组网，商业火箭发射频次提升。"
            "人形机器人：从'实验室Demo'到'工厂试点'的关键一年——"
            "Unitree/Figure/Tesla Bot谁能先在制造业场景跑通，谁就是下一个巨头。"
            "量子计算：量子比特数持续增长，但'量子优越性'的商业价值仍未兑现。"
        ),
        "mid": (
            "1-2年：人形机器人在制造业/物流/养老等场景的试点结果将决定这个赛道是'真风口'还是'泡沫'。"
            "可控核聚变的私人投资热潮可能降温——从'总能突破'到'商业化时间表仍不确定'的预期修正。"
        ),
        "long": (
            "3-5年：航天产业从'国家主导'向'国家+商业双轮驱动'转型——"
            "低轨卫星互联网、太空旅游、空间站商业化将创造全新的产业生态。"
            "脑机接口/可控核聚变/量子计算——这些'改变游戏规则'的技术可能在2030年前后迎来真正的拐点。"
        ),
        "confidence": "低-中",
    })

    forecasts.append({
        "sector": "投融资 / 创投",
        "hype_cycle": "泡沫低谷→稳步爬升过渡",
        "short": (
            "一级市场：募资难、退出难的问题短期内不会改变。"
            "资金集中押注AI赛道（尤其是大模型/Agent基础设施），非AI赛道融资难度加大。"
            "估值体系正在重构——'故事估值'退潮，'收入验证+AI溢价'成为新标准。"
        ),
        "mid": (
            "1-2年：IPO市场回暖的关键取决于两个变量——"
            "美联储降息路径的确定性和中美科技博弈的烈度。"
            "退出渠道多元化（港股/科创板/并购/分红退出）是VC行业的生存之道。"
        ),
        "long": "3-5年：创投行业本身正在被AI改造——AI辅助的deal sourcing和尽调将大幅提升早期投资的命中率。",
        "confidence": "中",
    })

    return forecasts


def analyze_emerging_opportunities(cats):
    """Identify white-space opportunities with low current coverage but high potential."""
    opportunities = []

    # Check what's missing from coverage
    covered_topics = set()
    for cat_items in cats.values():
        for item in cat_items:
            covered_topics.add(item.get("title", ""))

    # AI application verticals
    ai_vertical_kws = ["医疗AI", "AI制药", "AI教育", "AI法律", "AI金融", "工业AI", "AI for Science",
                        "蛋白质", "材料基因组", "AI农业"]
    ai_vertical_hits = sum(1 for t in covered_topics if any(k in t for k in ai_vertical_kws))
    if ai_vertical_hits < 3:
        opportunities.append({
            "field": "AI垂直行业应用（医疗/法律/教育/农业）",
            "why": "当前AI报道集中于基础模型竞争，AI在垂直行业的落地报道严重不足。这恰恰是价值创造最大的地带——'AI改造一个行业'比'做一个更好的模型'更具商业价值。",
            "signals_to_watch": "医疗AI的FDA/药监局审批进度、AI法律工具的律所采用率、工业AI的良品率提升数据",
            "time_to_bloom": "6-18个月",
        })

    # Quantum computing middleware
    quantum_kws = ["量子", "量子计算", "量子比特"]
    quantum_hits = sum(1 for t in covered_topics if any(k in t for k in quantum_kws))
    if quantum_hits < 2:
        opportunities.append({
            "field": "量子计算中间件与软件生态",
            "why": "量子计算的硬件竞赛吸引了大部分注意力，但量子算法/量子编译器/量子纠错软件/量子-经典混合架构这些'软件层'才是释放量子计算商业价值的关键。这是一个几乎空白的市场。",
            "signals_to_watch": "IBM Qiskit/Google Cirq/Pennylane等量子SDK的下载量增速、量子算法在金融/制药领域的论文数量",
            "time_to_bloom": "2-5年",
        })

    # Biotech + AI crossover
    bio_kws = ["生物", "基因", "合成生物", "mRNA", "细胞", "基因编辑", "CRISPR", "蛋白质"]
    bio_hits = sum(1 for t in covered_topics if any(k in t for k in bio_kws))
    if bio_hits < 3:
        opportunities.append({
            "field": "生物科技 × AI交叉（AI制药/合成生物学/精准医疗）",
            "why": "生物科技是AI之外最具变革性的技术领域。AI大幅降低了蛋白质结构预测和药物分子设计的成本——DeepMind的AlphaFold已证明了AI+生物的爆发力。当前媒体关注度远低于其实际潜力。",
            "signals_to_watch": "AI设计药物进入临床试验的数量、合成生物学公司的营收增速、FDA对AI辅助诊断的监管框架更新",
            "time_to_bloom": "1-5年",
        })

    # Smart manufacturing
    mfg_kws = ["智能制造", "工业互联网", "数字孪生", "工业4.0", "灯塔工厂"]
    mfg_hits = sum(1 for t in covered_topics if any(k in t for k in mfg_kws))
    if mfg_hits < 2:
        opportunities.append({
            "field": "智能制造与工业数字孪生",
            "why": "中国拥有全球最大规模的制造业，但智能制造的渗透率仍低。数字孪生+工业AI的组合可以在不新建工厂的情况下大幅提升现有产能的效率和柔性——这是'存量升级'的最大价值洼地。",
            "signals_to_watch": "工信部智能制造试点数量、工业互联网平台连接的设备数、灯塔工厂的中国新增数量",
            "time_to_bloom": "6-36个月",
        })

    return opportunities


def analyze_disruption_risks(cats, company_mentions):
    """Risk matrix: what could disrupt each sector and key companies."""
    risks = []

    top_companies = [c[0] for c in company_mentions[:8]]

    risks.append({
        "target": "AI/大模型赛道",
        "risk_type": "技术路线颠覆",
        "scenario": (
            "当前以Transformer为核心的技术路线可能被新架构（如状态空间模型Mamba/液态神经网络/"
            "神经符号AI）部分替代。如果新架构在推理效率和可解释性上取得突破，"
            "当前基于Transformer的巨额投入可能面临'资产搁浅'风险。"
        ),
        "probability": "中低",
        "impact": "极高",
        "early_signals": "Mamba/RWKV等新架构的论文引用增速、顶级AI会议的最佳论文方向变化",
    })

    risks.append({
        "target": "半导体产业链",
        "risk_type": "地缘政治断供",
        "scenario": (
            "如果台海局势升级或美国进一步扩大对华半导体出口管制范围（从先进制程扩展到成熟制程），"
            "全球半导体供应链将面临比2020-2022年更严重的断裂。"
            "考虑到全球75%的芯片产能集中在东亚，这种断裂的冲击将是系统性的。"
        ),
        "probability": "中",
        "impact": "灾难性",
        "early_signals": "美国BIS出口管制清单更新频率、台积电/三星的美国本土建厂进度、中国稀土出口管制动向",
    })

    risks.append({
        "target": "新能源/电动车",
        "risk_type": "产能过剩+价格战螺旋",
        "scenario": (
            "中国新能源车市场已进入'去产能'阵痛期——2026年产能利用率可能跌破60%。"
            "价格战从整车向上游电池/材料传导，整个产业链利润率承压。"
            "同时，欧美'去风险'政策（欧盟反补贴税、美国IRA本土化要求）限制了中国新能源企业的出海空间。"
        ),
        "probability": "高",
        "impact": "中高",
        "early_signals": "新能源车终端折扣率、动力电池价格走势、欧盟对华电动车的反补贴调查进展",
    })

    risks.append({
        "target": "互联网平台",
        "risk_type": "AI原生应用颠覆",
        "scenario": (
            "传统搜索/电商/社交可能被AI原生产品结构性替代——"
            "就像移动互联网颠覆了PC互联网的霸主，AI-Native产品有机会重构用户入口。"
            "搜索这个最赚钱的互联网商业模式首当其冲——ChatGPT/Perplexity式的'答案引擎'正在吃掉传统搜索的使用场景。"
        ),
        "probability": "中",
        "impact": "高",
        "early_signals": "Google搜索份额变化、Perplexity等AI搜索的用户增速、传统电商的流量成本变化趋势",
    })

    risks.append({
        "target": "头部科技企业",
        "risk_type": "估值修正风险",
        "scenario": (
            f"当前关注度最高的企业（{'、'.join(top_companies[:4])}等）的估值中包含了大量'AI溢价'。"
            f"如果AI的商业化落地速度低于市场预期（典型表现为企业AI收入占比迟迟不能突破10%），"
            f"这些溢价将面临修正。从历史规律看，技术泡沫破裂后头部公司的估值平均回调40-60%。"
        ),
        "probability": "中",
        "impact": "高",
        "early_signals": "头部科技公司的AI业务收入占比变化、AI相关CAPEX的回收周期、分析师对AI盈利时间表的调整",
    })

    return risks


def analyze_cross_domain_convergence(cats):
    """Cross-domain technology convergence predictions."""
    convergences = []

    convergences.append({
        "cross": "AI × 芯片 → 推理芯片 + 端侧AI芯片",
        "analysis": (
            "大模型从训练时代进入推理时代，芯片需求结构发生根本变化——"
            "训练芯片追求极致算力（英伟达H100/B200的逻辑），推理芯片追求能效比和低延迟。"
            "这为芯片创业公司打开了机会窗口——不必在训练芯片上和英伟达正面竞争，"
            "而是聚焦推理优化（Groq/LPUs/D-Matrix等新架构）。"
            "同时，端侧AI（手机/PC/汽车/可穿戴）将催生一个百亿级的NPU/APU芯片市场。"
        ),
        "timeline": "推理芯片爆发2026-2027，端侧AI芯片规模化2027-2028",
    })

    convergences.append({
        "cross": "AI × 机器人 → 具身智能产业化",
        "analysis": (
            "大模型给机器人装上了'大脑'——从过去只能执行预编程任务，进化到可以理解自然语言指令、"
            "泛化到新场景、从演示中学习。这一突破可能比大多数人想象的更快改变制造业和物流业。"
            "关键瓶颈不在AI（进步很快），而在硬件——灵巧手、关节电机、电池续航、成本控制。"
            "中国在硬件供应链上的优势+美国在AI算法上的优势，决定这个赛道的全球竞争格局。"
        ),
        "timeline": "2026-2027 工厂试点 → 2028-2029 仓储/物流规模化 → 2030+ 家庭服务机器人萌芽",
    })

    convergences.append({
        "cross": "新能源 × AI → 智能电网 + 材料加速发现",
        "analysis": (
            "AI正在加速新能源的两个关键环节："
            "1）AI驱动的电池材料发现——用机器学习筛选电极/电解液材料组合，将材料研发周期从10年缩短到2-3年。"
            "2）智能电网调度——新能源发电的间歇性和分布性需要极其复杂的电网调度算法，"
            "这正是深度强化学习的强项（DeepMind已证明AI可将数据中心能耗降低30%，同样的技术可用于电网）。"
            "这两个应用都不需要AGI——现有AI技术已经够用，缺的是工程化和行业know-how。"
        ),
        "timeline": "AI驱动材料发现已在进行，智能电网AI调度2027-2030逐步落地",
    })

    convergences.append({
        "cross": "航天 × 互联网 → 星地融合网络",
        "analysis": (
            "低轨卫星星座（Starlink/中国星网）将在2027年前后实现对全球无死角的覆盖。"
            "这不仅是'偏远地区上网'的问题——它将重塑全球通信产业的格局："
            "手机直连卫星（华为Mate 60已实现紧急通信→未来将普及到普通数据通信）、"
            "自动驾驶的高精度定位、全球IoT设备的实时连接。"
            "中国星网计划发射1.3万颗低轨卫星——这是对标Starlink的战略级布局。"
            "轨道资源和频率资源是有限的——先占先得，这个窗口不会永远敞开。"
        ),
        "timeline": "2026-2028 密集组网 → 2028-2030 商业运营规模化",
    })

    return convergences


def generate_technology_timeline(cats):
    """Structured technology milestone timeline."""
    milestones = []

    milestones.append({
        "phase": "2026 H2",
        "milestones": [
            ("GPT-5/Claude 4等级模型发布", "高", "训练算力规模和预训练数据量的进一步提升"),
            ("首批AI Agent产品在生产环境中规模化落地", "中高", "企业AI支出从实验性转向生产性"),
            ("固态电池量产验证结果揭晓", "中", "宁德时代/丰田/QuantumScape的量产良率数据"),
            ("中国商业火箭高频次发射", "高", "发射许可证发放频率和成功率"),
            ("人形机器人进入工厂试点", "中", "Figure/Tesla Bot/Unitree的试点客户公告"),
        ],
    })

    milestones.append({
        "phase": "2027",
        "milestones": [
            ("AI推理芯片市场格局初定", "中", "Groq/Cerebras/D-Matrix等新玩家的市场份额数据"),
            ("低轨卫星互联网开始商用服务", "中高", "Starlink/中国星网的用户数增长"),
            ("AI开始实质性影响白领就业市场", "中", "知识工作者的生产效率和薪资变化数据"),
            ("新能源车行业完成第一轮洗牌", "高", "尾部车企的退出/被并购数量"),
            ("量子计算达到1000+逻辑量子比特", "低-中", "IBM/Google/Quantinuum的量子体积基准测试"),
        ],
    })

    milestones.append({
        "phase": "2028+",
        "milestones": [
            ("AI跨越'早期大众'进入'晚期大众'阶段", "中", "企业AI采用率超过50%"),
            ("人形机器人在仓储/物流规模化部署", "低-中", "物流行业的机器人替代率数据"),
            ("固态电池电动车开始大规模交付", "中", "搭载固态电池的车型数量及销量"),
            ("硅基芯片逼近1nm物理极限", "高", "台积电/三星/Intel的工艺路线图更新"),
            ("可控核聚变实现Q>1的持续能量增益", "低", "ITER/私营聚变公司（Helion/CFS）的实验结果"),
        ],
    })

    return milestones


def generate_tech_outlook(cats, company_mentions, trends):
    """Comprehensive tech outlook — kept for backward compat, delegates to new functions."""
    forecasts = generate_sector_forecasts(cats)
    opportunities = analyze_emerging_opportunities(cats)
    risks = analyze_disruption_risks(cats, company_mentions)
    convergences = analyze_cross_domain_convergence(cats)
    milestones = generate_technology_timeline(cats)
    return {
        "forecasts": forecasts,
        "opportunities": opportunities,
        "risks": risks,
        "convergences": convergences,
        "milestones": milestones,
    }

# ═══════════════════ CHARTS ═══════════════════

def chart_tech_categories(cats):
    fig, ax = plt.subplots(figsize=(8, 4.5))
    ordered = ["AI/大模型", "半导体/芯片", "新能源/电动车", "互联网/平台", "航天/硬科技", "政策/监管", "投融资/创投"]
    counts = [len(cats.get(c, [])) for c in ordered]
    colors = [C_BLUE, C_RED, C_GREEN, C_ORANGE, C_PURPLE, C_GRAY, C_TEAL]
    bars = ax.bar(ordered, counts, color=colors, edgecolor="white", linewidth=0.5)
    for bar, v in zip(bars, counts):
        if v > 0:
            ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.3, str(v),
                    ha="center", fontsize=9, fontweight="bold")
    ax.set_title("科技产业新闻分布", fontsize=14, fontweight="bold", pad=12)
    ax.set_ylabel("新闻数量")
    plt.xticks(rotation=30, ha="right", fontsize=8)
    ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
    fig.tight_layout()
    path = f"{CHART_DIR}/tech_cats.png"
    fig.savefig(path, dpi=150, bbox_inches="tight"); plt.close(fig)
    return path

def chart_company_mentions(mentions):
    fig, ax = plt.subplots(figsize=(7, 4))
    names = [m[0] for m in mentions[:10]][::-1]
    counts = [m[1] for m in mentions[:10]][::-1]
    colors = plt.cm.Blues([0.3 + 0.7 * i/len(names) for i in range(len(names))])
    ax.barh(names, counts, color=colors, edgecolor="white")
    for i, v in enumerate(counts):
        ax.text(v + 0.3, i, str(v), va="center", fontsize=9, fontweight="bold")
    ax.set_title("科技企业媒体关注度 TOP 10", fontsize=14, fontweight="bold", pad=12)
    ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
    fig.tight_layout()
    path = f"{CHART_DIR}/company_mentions.png"
    fig.savefig(path, dpi=150, bbox_inches="tight"); plt.close(fig)
    return path

def chart_trend_timeline(cats, title_str, filename):
    all_dates = []
    for items in cats.values():
        for item in items:
            ctime = item.get("ctime", "")
            if ctime:
                try:
                    ts = int(ctime)
                    all_dates.append(datetime.fromtimestamp(ts).strftime("%m-%d"))
                except: pass
    if not all_dates: return None
    date_counts = Counter(all_dates)
    dates = sorted(date_counts.keys())[-14:]
    counts = [date_counts[d] for d in dates]
    fig, ax = plt.subplots(figsize=(8, 3))
    ax.fill_between(range(len(dates)), counts, alpha=0.3, color=C_BLUE)
    ax.plot(range(len(dates)), counts, marker="o", color=C_BLUE, linewidth=2, markersize=4)
    ax.set_xticks(range(len(dates)))
    ax.set_xticklabels(dates, rotation=45, fontsize=7)
    ax.set_title(title_str, fontsize=13, fontweight="bold", pad=12)
    ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
    fig.tight_layout()
    path = f"{CHART_DIR}/{filename}.png"
    fig.savefig(path, dpi=150, bbox_inches="tight"); plt.close(fig)
    return path

def chart_tech_roadmap(milestones):
    """Technology roadmap: timeline with maturity stages."""
    fig, ax = plt.subplots(figsize=(10, 5.5))

    phases = [m["phase"] for m in milestones]
    all_items = []
    phase_boundaries = [0]
    for m in milestones:
        items = m["milestones"]
        all_items.extend(items)
        phase_boundaries.append(len(all_items))

    phase_colors = {"2026 H2": C_BLUE, "2027": C_ORANGE, "2028+": C_GREEN}
    y_positions = list(range(len(all_items)))
    y_labels = [f"{item[0]}  [{item[1]}]" for item in all_items]

    colors = []
    current_phase_idx = 0
    for i in range(len(all_items)):
        while current_phase_idx < len(phase_boundaries) - 1 and i >= phase_boundaries[current_phase_idx + 1]:
            current_phase_idx += 1
        colors.append(phase_colors.get(phases[current_phase_idx], C_GRAY))

    for phase_idx, phase in enumerate(phases):
        start_y = -0.5 if phase_idx == 0 else phase_boundaries[phase_idx] - 0.5
        end_y = len(all_items) - 0.5
        if phase_idx < len(phases) - 1:
            end_y = phase_boundaries[phase_idx + 1] - 0.5
        ax.axhspan(start_y, end_y, alpha=0.06, color=phase_colors[phase])
        ax.text(0.98, (start_y + end_y) / 2, phase,
                transform=ax.get_yaxis_transform(), ha="right", va="center",
                fontsize=11, fontweight="bold", color=phase_colors[phase], alpha=0.6)

    ax.barh(y_labels, [0.5] * len(all_items), color=colors, edgecolor="white", height=0.7)
    ax.set_xlim(0, 1)
    ax.axis("off")

    for i, (label, confidence, signal) in enumerate(all_items):
        ax.text(0.52, i, f"领先指标: {signal}", va="center", fontsize=7, color=C_GRAY)

    ax.set_title("技术里程碑路线图", fontsize=14, fontweight="bold", pad=16)
    legend_elements = [plt.Rectangle((0, 0), 1, 1, facecolor=phase_colors[p], alpha=0.3, label=p) for p in phases]
    ax.legend(handles=legend_elements, loc="lower right", fontsize=9, ncol=3)

    fig.tight_layout()
    path = f"{CHART_DIR}/tech_roadmap.png"
    fig.savefig(path, dpi=150, bbox_inches="tight")
    plt.close(fig)
    return path


# ═══════════════════ MAIN ═══════════════════

print("=" * 60)
print("科技产业分析日报生成器 v2.0 — 前瞻性分析")
print("=" * 60)

print("\n[1/4] 获取科技新闻...")
tech_keywords = ["人工智能", "芯片半导体", "新能源电动车", "航天卫星", "互联网科技",
                  "AI大模型", "机器人", "量子计算", "生物科技", "科技政策"]
tech_raw = fetch_sina_search(tech_keywords, size=20)
tech_raw += fetch_sina_news(lid=2509, num=40)
seen = set()
tech_news = []
for item in tech_raw:
    u = item.get("url", "")
    if u not in seen:
        seen.add(u)
        tech_news.append(item)
print(f"  科技新闻: {len(tech_news)} 条")

print("\n[2/4] 执行分析...")
tech_cats = classify_tech(tech_news)
company_mentions = analyze_company_mentions(tech_news)
tech_trends = analyze_tech_trends(tech_cats, company_mentions)
tech_outlook = generate_tech_outlook(tech_cats, company_mentions, tech_trends)
for k, v in tech_cats.items():
    if v: print(f"  {k}: {len(v)} 条")
print(f"  企业关注度 TOP 5: {', '.join(f'{c}({n})' for c,n in company_mentions[:5])}")
print(f"  前瞻分析: {len(tech_outlook['forecasts'])} 赛道预测 | {len(tech_outlook['opportunities'])} 新兴机会 | {len(tech_outlook['risks'])} 风险场景 | {sum(len(m['milestones']) for m in tech_outlook['milestones'])} 技术里程碑")

print("\n[3/4] 生成图表...")
charts = {
    "tech_cats": chart_tech_categories(tech_cats),
    "company_mentions": chart_company_mentions(company_mentions),
}
tl = chart_trend_timeline(tech_cats, "科技新闻热度趋势", "trend_timeline")
if tl: charts["timeline"] = tl
charts["roadmap"] = chart_tech_roadmap(tech_outlook["milestones"])
print(f"  {len(charts)} 张图表已生成")

# ═══════════════════ BUILD DOCX ═══════════════════

print("\n[4/4] 生成报告文档...")

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2); section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'; style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def fmt_date(ctime_str):
    try:
        ts = int(ctime_str)
        return datetime.fromtimestamp(ts).strftime("%Y-%m-%d")
    except: return ""

def add_bullet(doc, item):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(item["title"]); run.font.size = Pt(10)
    d = fmt_date(item.get("ctime", ""))
    if d:
        run2 = p.add_run(f'  ({d})'); run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(150, 150, 150)

def add_analysis(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f'▸ 分析：{text}'); run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(int(C_BLUE[1:3],16), int(C_BLUE[3:5],16), int(C_BLUE[5:7],16))
    p.paragraph_format.left_indent = Cm(0.8)

# COVER
title = doc.add_heading('科技产业分析日报', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph(); subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run(f'{datetime.now().strftime("%Y年%m月%d日")} | 科技产业深度分析')
run.font.size = Pt(12); run.font.color.rgb = RGBColor(100, 100, 100)
doc.add_paragraph()

# Executive Summary
doc.add_heading('核心发现', level=1)
total = sum(len(v) for v in tech_cats.values())
ai_n = len(tech_cats.get("AI/大模型", [])); chip_n = len(tech_cats.get("半导体/芯片", []))
ev_n = len(tech_cats.get("新能源/电动车", []))
doc.add_paragraph(
    f"本日监测科技产业新闻{total}条，覆盖{sum(1 for v in tech_cats.values() if v)}个细分赛道。"
    f"AI/大模型（{ai_n}条）、半导体/芯片（{chip_n}条）、新能源/电动车（{ev_n}条）为三大热点。"
    f"企业关注度前三：{'、'.join(f'{c}({n}次)' for c,n in company_mentions[:3])}。"
)

# ── SECTION MAP ──
sections = [
    ("一、AI / 大模型", "AI/大模型",
     "从技术采用生命周期看，AI正在跨越早期采用者到早期大众的'鸿沟'。当前阶段的核心矛盾不是模型能力不够，而是企业落地成本太高、人才供给不足、以及监管不确定性。基础设施层（算力、数据、工具链）的投资确定性高于应用层。"),
    ("二、半导体 / 芯片", "半导体/芯片",
     "半导体是科技竞争的'底层棋局'。美国主导的设备/EDA/先进制程封锁与中国'举国体制'式自主替代之间的拉锯，决定了整个科技产业的上游供应稳定性。从产业链安全视角看，成熟制程（28nm及以上）的国产化率快速提升，先进制程（7nm以下）仍是瓶颈。"),
    ("三、新能源 / 电动车", "新能源/电动车",
     "新能源产业已从'政策驱动'进入'市场驱动+技术驱动'双轮阶段。固态电池量产时间表、智能化水平、海外市场准入三者共同决定下一轮行业洗牌的节奏。中国企业在电池、光伏、风电领域的全球份额领先，但面临欧美'去风险'政策压力。"),
    ("四、互联网 / 平台经济", "互联网/平台",
     "平台经济监管进入常态化阶段。增长引擎从'流量红利'转向'AI赋能+出海'。字节/TikTok的全球化、 Temu/SHEIN的跨境电商、以及AI对搜索/推荐/广告的重构，是当前最值得关注的三个结构性变量。"),
    ("五、航天 / 硬科技", "航天/硬科技",
     "商业航天、量子计算、可控核聚变、脑机接口、人形机器人——这些'硬科技'赛道虽短期商业化前景不明，但代表了科技产业的长期方向。从风险投资视角看，这些赛道适合'小额度、多标的、长周期'的组合策略。"),
    ("六、科技政策与监管", "政策/监管",
     "科技政策的底层逻辑是'在发展和安全之间寻找动态平衡'。AI立法、数据跨境流动、反垄断执法、芯片出口管制是四个政策主线。政策变化的领先指标包括：人大常委会立法计划、国务院产业规划、以及网信办/工信部的部门规章更新。"),
    ("七、投融资与创投动态", "投融资/创投",
     "一级市场投融资数据是科技产业未来2-3年的'前置指标'。当前全球科技投资从'撒胡椒面'模式转向'集中押注'模式——资金向头部项目和确定性赛道集中，腰部项目融资难度加大。"),
]

for title, cat_key, analysis_text in sections:
    doc.add_heading(title, level=2)
    items = tech_cats.get(cat_key, [])
    if items:
        for item in items[:8]: add_bullet(doc, item)
    else:
        doc.add_paragraph("本日该领域无显著新闻。")
    if items:
        doc.add_paragraph()
        add_analysis(doc, f"{analysis_text} 本日该赛道新闻{len(items)}条，{'信息密度较高，值得重点关注。' if len(items) > 5 else '保持常规关注。'}")

# Charts
doc.add_paragraph()
if charts.get("tech_cats"):
    doc.add_paragraph().add_run().add_picture(charts["tech_cats"], width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
if charts.get("company_mentions"):
    doc.add_paragraph().add_run().add_picture(charts["company_mentions"], width=Inches(5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Tech Trends Deep Dive
doc.add_paragraph()
doc.add_heading('八、技术趋势深度研判', level=2)
for trend in tech_trends:
    p_t = doc.add_paragraph()
    run = p_t.add_run(trend["title"]); run.bold = True; run.font.size = Pt(12)
    p_m = doc.add_paragraph()
    run_m = p_m.add_run(f'信号: {trend["signal"]}　|　时间维度: {trend["horizon"]}')
    run_m.font.size = Pt(10); run_m.font.color.rgb = RGBColor(int(C_RED[1:3],16), int(C_RED[3:5],16), int(C_RED[5:7],16))
    doc.add_paragraph(trend["analysis"])

# ── Cross-domain convergence ──
doc.add_paragraph()
doc.add_heading('八-A、跨域融合前瞻', level=2)
add_analysis(doc, "技术的最大突破往往发生在学科交叉地带。以下四个交叉领域正在酝酿下一波「大机会」。")
doc.add_paragraph()
for cv in tech_outlook["convergences"]:
    p_t = doc.add_paragraph()
    run = p_t.add_run(f'⚙ {cv["cross"]}'); run.bold = True; run.font.size = Pt(12)
    doc.add_paragraph(cv["analysis"])
    p_tl = doc.add_paragraph()
    run_tl = p_tl.add_run(f'时间轴: {cv["timeline"]}'); run_tl.font.size = Pt(10)
    run_tl.font.color.rgb = RGBColor(int(C_GREEN[1:3],16), int(C_GREEN[3:5],16), int(C_GREEN[5:7],16))

# ═══════════════════ OUTLOOK: 前瞻性分析 ═══════════════════

doc.add_paragraph()
doc.add_heading('九、科技产业前瞻', level=1)

# 9.1 Technology Roadmap
doc.add_heading('9.1 技术里程碑路线图', level=2)
add_analysis(doc, "以下是基于当前技术进展和政策信号推断的关键里程碑。置信度基于历史规律、投入规模和技术成熟度的交叉评估。")
if charts.get("roadmap"):
    doc.add_paragraph()
    doc.add_paragraph().add_run().add_picture(charts["roadmap"], width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# 9.2 Sector Forecasts
doc.add_heading('9.2 分赛道预测', level=2)
for fc in tech_outlook["forecasts"]:
    # Sector header with Hype Cycle position
    p_sector = doc.add_paragraph()
    run_s = p_sector.add_run(fc["sector"]); run_s.bold = True; run_s.font.size = Pt(13)
    p_hc = doc.add_paragraph()
    run_hc = p_hc.add_run(f'Hype Cycle: {fc["hype_cycle"]}　|　预测置信度: {fc["confidence"]}')
    run_hc.font.size = Pt(9); run_hc.font.color.rgb = RGBColor(150, 150, 150)

    # Short-term
    p_st = doc.add_paragraph()
    run_st = p_st.add_run("6个月预测："); run_st.bold = True; run_st.font.size = Pt(11)
    run_st.font.color.rgb = RGBColor(int(C_RED[1:3],16), int(C_RED[3:5],16), int(C_RED[5:7],16))
    doc.add_paragraph(fc["short"])

    # Mid-term
    p_mt = doc.add_paragraph()
    run_mt = p_mt.add_run("1-2年预测："); run_mt.bold = True; run_mt.font.size = Pt(11)
    run_mt.font.color.rgb = RGBColor(int(C_ORANGE[1:3],16), int(C_ORANGE[3:5],16), int(C_ORANGE[5:7],16))
    doc.add_paragraph(fc["mid"])

    # Long-term
    p_lt = doc.add_paragraph()
    run_lt = p_lt.add_run("3-5年愿景："); run_lt.bold = True; run_lt.font.size = Pt(11)
    run_lt.font.color.rgb = RGBColor(int(C_BLUE[1:3],16), int(C_BLUE[3:5],16), int(C_BLUE[5:7],16))
    doc.add_paragraph(fc["long"])
    doc.add_paragraph()

# 9.3 Emerging Opportunities
if tech_outlook["opportunities"]:
    doc.add_heading('9.3 新兴机会地图', level=2)
    add_analysis(doc, "以下领域当前媒体关注度低于其实际潜力——'信息洼地'往往意味着'价值洼地'。")
    doc.add_paragraph()
    for op in tech_outlook["opportunities"]:
        p_f = doc.add_paragraph()
        run_f = p_f.add_run(f'💡 {op["field"]}'); run_f.bold = True; run_f.font.size = Pt(12)
        doc.add_paragraph(op["why"])
        p_sig = doc.add_paragraph()
        run_sig = p_sig.add_run(f'领先信号: {op["signals_to_watch"]}'); run_sig.font.size = Pt(10)
        p_time = doc.add_paragraph()
        run_time = p_time.add_run(f'预计窗口期: {op["time_to_bloom"]}'); run_time.font.size = Pt(10)
        run_time.font.color.rgb = RGBColor(int(C_TEAL[1:3],16), int(C_TEAL[3:5],16), int(C_TEAL[5:7],16))
        doc.add_paragraph()

# 9.4 Disruption Risks
if tech_outlook["risks"]:
    doc.add_heading('9.4 颠覆风险矩阵', level=2)
    add_analysis(doc, "前瞻性分析不能只谈机会不谈风险。以下是当前构成最大威胁的颠覆场景——小概率、大破坏的事件尤其值得警惕。")
    doc.add_paragraph()
    for risk in tech_outlook["risks"]:
        p_r = doc.add_paragraph()
        run_r = p_r.add_run(f'⚠ {risk["target"]} — {risk["risk_type"]}'); run_r.bold = True; run_r.font.size = Pt(12)
        doc.add_paragraph(risk["scenario"])
        p_meta = doc.add_paragraph()
        run_meta = p_meta.add_run(f'概率: {risk["probability"]}　|　影响: {risk["impact"]}　|　领先信号: {risk["early_signals"]}')
        run_meta.font.size = Pt(9); run_meta.font.color.rgb = RGBColor(150, 150, 150)
        doc.add_paragraph()

# 9.5 Synthesis
doc.add_heading('9.5 综合研判', level=2)
forecast_count = len(tech_outlook["forecasts"])
opp_count = len(tech_outlook["opportunities"])
risk_count = len(tech_outlook["risks"])
convergence_count = len(tech_outlook["convergences"])

synthesis_text = (
    f"基于对{forecast_count}个赛道的分阶段预测、{opp_count}个新兴机会的识别、{risk_count}个颠覆风险的评估、"
    f"以及{convergence_count}个跨域融合方向的分析，当前科技产业的前瞻性判断如下：\n\n"
    f"核心主线：AI的推理成本持续下降是未来3年最确定的结构性趋势——它将像电力成本下降一样"
    f"渗透到每一个行业。抓住'AI使能'逻辑比押注'哪个模型最强'更重要。\n\n"
    f"最大不确定性：中美科技博弈的烈度和走向。它决定了半导体供应链的结构、AI人才的流向、"
    f"以及中国科技企业出海的天花板。这个变量不是'技术问题'，但它是所有技术预测的最大扰动项。\n\n"
    f"最被低估的方向：AI+生物科技和AI+工业制造的交叉领域。这两个方向当前媒体关注度远低于大模型竞赛，"
    f"但一旦突破，其对GDP的乘数效应将远大于'更好的聊天机器人'。\n\n"
    f"行动框架：从投资/职业/创业视角看，(1)短期（6-12个月）关注AI Agent和推理芯片的落地进展；"
    f"(2)中期（1-3年）布局端侧AI和智能制造的渗透率提升；"
    f"(3)长期（3-5年）跟踪量子计算、可控核聚变和人形机器人的商业化拐点。"
)
doc.add_paragraph(synthesis_text)

# FOOTER
doc.add_paragraph(); doc.add_paragraph()
footer = doc.add_paragraph(); footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('— 以上分析基于公开信息，纯属研究性质，不构成投资建议 —')
run.font.size = Pt(9); run.font.color.rgb = RGBColor(150, 150, 150); run.italic = True
footer2 = doc.add_paragraph(); footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer2.add_run(
    f'数据: 新浪新闻 | 框架: Hype Cycle / 波特五力 / 技术采用生命周期 '
    f'| 生成: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
run.font.size = Pt(8); run.font.color.rgb = RGBColor(150, 150, 150)

# SAVE
output_dir = '/Users/wwwqwq/Desktop/项目/科技报告'
os.makedirs(output_dir, exist_ok=True)
output_path = f'{output_dir}/科技产业分析_{datetime.now().strftime("%Y-%m-%d")}.docx'
doc.save(output_path)
print(f'\n{"="*60}')
print(f'报告已生成: {output_path}')
print(f'科技新闻: {total} 条 | {sum(1 for v in tech_cats.values() if v)} 个赛道')
print(f'含 {len(charts)} 张图表')
print(f'{"="*60}')
